from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
SOURCE = HERE / "report.md"
OUTPUT = HERE / "metacognition_gemini_3_1_flash_lite_preview.docx"

ACCENT = RGBColor(37, 79, 112)
MUTED = RGBColor(97, 107, 117)
BORDER = "D9DEE4"
HEADER_FILL = "EAF1F7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_twips: int) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    title = styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = ACCENT

    for name, size, before, after in (
        ("Heading 1", 15, 12, 5),
        ("Heading 2", 13, 10, 4),
        ("Heading 3", 11, 8, 3),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "AutoPsych metacognition report"
    header.style = styles["Normal"]
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def split_markdown_table_row(line: str) -> list[str]:
    row = line.strip().strip("|")
    return [cell.strip() for cell in row.split("|")]


def is_separator_row(line: str) -> bool:
    return bool(re.match(r"^\|\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$", line.strip()))


def clean_inline(text: str) -> str:
    text = re.sub(r"!\[[^\]]*\]\([^)]+\)", "", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("`", "")
    return text


def add_inline_text(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part[1:-1] if part.startswith("`") and part.endswith("`") else part)
        run.font.name = "Courier New" if part.startswith("`") and part.endswith("`") else "Arial"
        if part.startswith("`") and part.endswith("`"):
            run.font.size = Pt(9.5)


def add_markdown_table(document: Document, lines: list[str]) -> None:
    header = split_markdown_table_row(lines[0])
    body = [split_markdown_table_row(line) for line in lines[2:] if line.strip()]
    col_count = len(header)
    table = document.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"
    table.autofit = False
    set_table_width(table, 13680)

    wide_columns = {0}
    if col_count >= 6:
        widths = [2400] + [1700] * (col_count - 1)
    elif col_count == 3:
        widths = [2600, 2400, 2600]
    else:
        widths = [int(13680 / col_count)] * col_count

    for idx, value in enumerate(header):
        cell = table.rows[0].cells[idx]
        cell.text = clean_inline(value)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.width = Pt(widths[idx] / 20)
        set_cell_shading(cell, HEADER_FILL)
        set_cell_borders(cell)
        set_cell_margins(cell)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx in wide_columns else WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.bold = True
                run.font.size = Pt(8.5)
                run.font.color.rgb = ACCENT

    for body_row in body:
        cells = table.add_row().cells
        for idx, value in enumerate(body_row[:col_count]):
            cell = cells[idx]
            cell.text = clean_inline(value)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = Pt(widths[idx] / 20)
            set_cell_borders(cell)
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx in wide_columns else WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8.2 if col_count >= 6 else 9.2)

    document.add_paragraph()


def add_image(document: Document, line: str) -> None:
    match = re.search(r"!\[([^\]]*)\]\(([^)]+)\)", line)
    if not match:
        return
    caption, target = match.groups()
    image_path = (HERE / target).resolve()
    if image_path.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(7.8))
        if caption:
            cap = document.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = MUTED


def convert() -> None:
    document = Document()
    configure_document(document)

    lines = SOURCE.read_text().splitlines()
    idx = 0
    while idx < len(lines):
        line = lines[idx].rstrip()
        if not line:
            idx += 1
            continue

        if line.startswith("# "):
            paragraph = document.add_paragraph(style="Title")
            add_inline_text(paragraph, line[2:].strip())
        elif line.startswith("## "):
            document.add_paragraph(clean_inline(line[3:].strip()), style="Heading 1")
        elif line.startswith("### "):
            document.add_paragraph(clean_inline(line[4:].strip()), style="Heading 2")
        elif line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_text(paragraph, line[2:].strip())
        elif line.startswith("!["):
            add_image(document, line)
        elif line.startswith("|") and idx + 1 < len(lines) and is_separator_row(lines[idx + 1]):
            table_lines = [line, lines[idx + 1].rstrip()]
            idx += 2
            while idx < len(lines) and lines[idx].startswith("|"):
                table_lines.append(lines[idx].rstrip())
                idx += 1
            add_markdown_table(document, table_lines)
            continue
        else:
            paragraph = document.add_paragraph()
            add_inline_text(paragraph, line)
        idx += 1

    document.core_properties.title = "Metacognition in google/gemini-3.1-flash-lite-preview"
    document.core_properties.subject = "AutoPsych experimental report"
    document.core_properties.author = "AutoPsych"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    convert()
